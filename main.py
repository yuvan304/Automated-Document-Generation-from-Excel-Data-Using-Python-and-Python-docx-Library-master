import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
from datetime import datetime

def update_document(excel_data, docx_template):
    for idx, row in excel_data.iterrows():
        if row['Standard_letter_ID'] == 'BT129':
            # Load the Word document
            doc = Document(docx_template)

            # 1. Update date
            for paragraph in doc.paragraphs:
                if 'Click here to enter a date.' in paragraph.text:
                    paragraph.text = paragraph.text.replace('Click here to enter a date.', datetime.now().strftime('%Y-%m-%d'))

            # 2. Update [Customer First Name Last Name]
            for paragraph in doc.paragraphs:
                if '[Customer First Name Last Name]' in paragraph.text:
                    paragraph.text = paragraph.text.replace('[Customer First Name Last Name]', row['PRIM_CUST_FULLNAME'])

            # 3. Update [Address] with state
            for paragraph in doc.paragraphs:
                if '[Address]' in paragraph.text:
                    paragraph.text = paragraph.text.replace('[Address]', row['state'])

            # 4. Update [Address] with zip
            for paragraph in doc.paragraphs:
                if '[Address]' in paragraph.text:
                    paragraph.text = paragraph.text.replace('[Address]', str(row['zip']))

            # 5. Update Account/Application Numbers dropdown
            for content_control in doc.inline_shapes:
                if content_control.type == 3:  # InlineShapeType.TEXT_BOX
                    textbox = content_control._inline.graphic.graphicData.docPr
                    if textbox.title == "Account/Application Numbers":
                        dropdown = content_control._element.find('.//w:comboBox')
                        dropdown.clear_content()
                        dropdownlist = OxmlElement('w:listEntry')
                        dropdownlist.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', f"xxxx{str(row['acc_num'])[-4:]}")
                        dropdownlist.text = "choose account number ending in:"
                        dropdown.append(dropdownlist)

            # 6. Update [First Name Last Name]
            for paragraph in doc.paragraphs:
                if '[First Name Last Name]' in paragraph.text:
                    paragraph.text = paragraph.text.replace('[First Name Last Name]', row['PRIM_CUST_FULLNAME'])

            # 7. Update Resolution Openings dropdown
            for content_control in doc.inline_shapes:
                if content_control.type == 3:  # InlineShapeType.TEXT_BOX
                    textbox = content_control._inline.graphic.graphicData.docPr
                    if textbox.title == "Resolution Openings":
                        dropdown = content_control._element.find('.//w:comboBox')
                        dropdown.clear_content()
                        dropdownlist = OxmlElement('w:listEntry')
                        dropdownlist.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', "Thank you for contacting us. We completed our research and are providing details to address your concerns.")
                        dropdownlist.text = "choose an item alias Resolution Openings"
                        dropdown.append(dropdownlist)

            # 8. Update Me/Us dropdown
            for content_control in doc.inline_shapes:
                if content_control.type == 3:  # InlineShapeType.TEXT_BOX
                    textbox = content_control._inline.graphic.graphicData.docPr
                    if textbox.title == "Me/Us":
                        dropdown = content_control._element.find('.//w:comboBox')
                        dropdown.clear_content()
                        dropdownlist = OxmlElement('w:listEntry')
                        dropdownlist.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', "us")
                        dropdownlist.text = "choose an item alias Me/Us"
                        dropdown.append(dropdownlist)

            # 9. Update contact Us dropdown
            for content_control in doc.inline_shapes:
                if content_control.type == 3:  # InlineShapeType.TEXT_BOX
                    textbox = content_control._inline.graphic.graphicData.docPr
                    if textbox.title == "contact Us":
                        dropdown = content_control._element.find('.//w:comboBox')
                        dropdown.clear_content()
                        dropdownlist = OxmlElement('w:listEntry')
                        dropdownlist.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', "reach us @291090")
                        dropdownlist.text = "choose an item alias contact Us"
                        dropdown.append(dropdownlist)

            # Save the updated document with the CASE_NR as filename in the output directory
            doc.save(f"C:\\Users\\arjun\\Downloads\\op\\{row['CASE_NR']}.docx")

# Load Excel data
excel_data = pd.read_excel(r'C:\Users\arjun\Downloads\sample.xlsx')

# Load Word document template
docx_template = r'C:\Users\arjun\Downloads\sample.docx'

# Call the function to update the document
update_document(excel_data, docx_template)
