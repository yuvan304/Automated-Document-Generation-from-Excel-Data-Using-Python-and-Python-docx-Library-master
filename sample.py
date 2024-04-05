import pandas as pd
from docx import Document
from docx.shared import Pt
from datetime import datetime

def update_document(excel_data, docx_template):
    for idx, row in excel_data.iterrows():
        if row['Standard_letter_ID'] == 'BT129':
            # Load the Word document
            doc = Document(docx_template)

            # Update date in the first cell of the table
            current_date = datetime.now().strftime('%Y-%m-%d')
            table = doc.tables[0]  # Assuming the table is the first table in the document
            cell = table.cell(0, 0)  # First row, first cell
            cell.text = current_date

            # Update [Customer First Name Last Name]

            cell=table.cell(1,0)
            for paragraph in doc.paragraphs:
               cell.text = (row['PRIM_CUST_FULLNAME'])

            # Update [Address] with state
            cell=table.cell(2,0)
            for paragraph in cell.paragraphs:
                paragraph.text = paragraph.text.replace('[State]', row['state'])

            # Update [Address] with zip
            # cell=table.cell(4,0)
            for paragraph in cell.paragraphs:
                paragraph.text = paragraph.text.replace('[zip]', str(row['zip']))

            # Update Account/Application Numbers dropdown

            for paragraph in doc.paragraphs:
                if 'xxxx' in paragraph.text:
                    paragraph.text = paragraph.text.replace('xxxx', str(row['acc_num'])[-4:])
            # Update [First Name Last Name]
            for paragraph in doc.paragraphs:
                if '[First Name Last Name]' in paragraph.text:
                    paragraph.text = paragraph.text.replace('[First Name Last Name]', row['PRIM_CUST_FULLNAME'])

            # Update Resolution Openings dropdown
            for content_control in doc.inline_shapes:
                if content_control.type == 3:  # InlineShapeType.TEXT_BOX
                    textbox = content_control._inline.graphic.graphicData.docPr
                    if textbox.title == "Resolution Openings":
                        # Update the dropdown value
                        dropdownlist = content_control._element.find('.//w:listEntry')
                        dropdownlist.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', "Thank you for contacting us. We completed our research and are providing details to address your concerns.")

            # Update Me/Us dropdown
            for content_control in doc.inline_shapes:
                if content_control.type == 3:  # InlineShapeType.TEXT_BOX
                    textbox = content_control._inline.graphic.graphicData.docPr
                    if textbox.title == "Me/Us":
                        # Update the dropdown value
                        dropdownlist = content_control._element.find('.//w:listEntry')
                        dropdownlist.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', "us")

            # Update contact Us dropdown
            for content_control in doc.inline_shapes:
                if content_control.type == 3:  # InlineShapeType.TEXT_BOX
                    textbox = content_control._inline.graphic.graphicData.docPr
                    if textbox.title == "contact Us":
                        # Update the dropdown value
                        dropdownlist = content_control._element.find('.//w:listEntry')
                        dropdownlist.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', "reach us @291090")

            # Save the updated document with the CASE_NR as filename in the output directory
            doc.save(f"C:\\Users\\arjun\\Downloads\\op\\{row['CASE_NR']}.docx")

# Load Excel data
excel_data = pd.read_excel(r'C:\Users\arjun\Downloads\sample.xlsx')

# Load Word document template
docx_template = r'C:\Users\arjun\\Downloads\sample.docx'

# Call the function to update the document
update_document(excel_data, docx_template)
